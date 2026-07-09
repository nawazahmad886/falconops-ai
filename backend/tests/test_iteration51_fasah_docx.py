"""
Iteration 51 — Fasah Weekly Report DOCX export end-to-end tests.

Covers:
- POST /api/events/upload  (Excel xlsx upload)
- POST /api/events/analyze/{upload_id}
- GET  /api/events/export/{analysis_id}/docx
- DOCX structural verification (python-docx)
- Severity color coding (w:shd w:fill XML)
- Branding params (company, period)
- Auth guard, 404 handling
- Existing /excel and /pdf still work
- Mixed-case header normalization
- Empty alert_frequency case (still 200, no crash)
- AI Event Analyzer related routes still functional
"""
import io
import os
import re
import zipfile
from datetime import datetime, timedelta

import pytest
import requests
from openpyxl import Workbook
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://health-rules-engine.preview.emergentagent.com").rstrip("/")
ADMIN_EMAIL = "admin@falconapps.com"
ADMIN_PASSWORD = "Admin@123"
VIEWER_EMAIL = "test@falconapps.com"
VIEWER_PASSWORD = "testpass123"

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ───────────────────────── fixtures ─────────────────────────

@pytest.fixture(scope="module")
def admin_token():
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"email": ADMIN_EMAIL, "password": ADMIN_PASSWORD},
                      timeout=30)
    assert r.status_code == 200, f"admin login failed: {r.status_code} {r.text}"
    tok = r.json().get("access_token")
    assert tok
    return tok


@pytest.fixture(scope="module")
def admin_headers(admin_token):
    return {"Authorization": f"Bearer {admin_token}"}


def _make_xlsx(rows, headers=("timestamp", "service", "alert", "severity")):
    wb = Workbook()
    ws = wb.active
    ws.append(list(headers))
    for r in rows:
        ws.append(list(r))
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _sample_rows():
    base = datetime(2026, 4, 12, 10, 0, 0)
    rows = []
    services = ["payment-api", "checkout-api", "search-api", "auth-api", "billing-api"]
    alerts = [
        ("Database connection timeout", "critical"),
        ("API response latency high", "warning"),
        ("CPU utilization high", "critical"),
        ("Memory leak detected", "warning"),
        ("Disk space low", "info"),
    ]
    # 13 rows mixing severities and services
    pattern = [
        (0, 0), (0, 0), (1, 1), (2, 2), (2, 2),
        (3, 3), (4, 4), (1, 1), (0, 0), (2, 2),
        (3, 3), (1, 1), (4, 4),
    ]
    for i, (svc_i, alt_i) in enumerate(pattern):
        ts = base + timedelta(minutes=5 * i)
        alert, sev = alerts[alt_i]
        rows.append([ts.strftime("%Y-%m-%d %H:%M:%S"), services[svc_i], alert, sev])
    return rows


@pytest.fixture(scope="module")
def analysis_id(admin_headers):
    """Upload + analyze an Excel; yield analysis_id and upload_id."""
    xlsx = _make_xlsx(_sample_rows())
    files = {"file": ("TEST_fasah_sample.xlsx", xlsx, "application/octet-stream")}
    r = requests.post(f"{BASE_URL}/api/events/upload",
                      headers=admin_headers, files=files, timeout=60)
    assert r.status_code == 200, f"upload failed: {r.status_code} {r.text}"
    upload_id = r.json()["upload_id"]

    r2 = requests.post(f"{BASE_URL}/api/events/analyze/{upload_id}",
                       headers=admin_headers, timeout=180)
    assert r2.status_code == 200, f"analyze failed: {r2.status_code} {r2.text}"
    aid = r2.json().get("analysis_id")
    assert aid
    return {"analysis_id": aid, "upload_id": upload_id}


# ─────────────────────── auth/health basics ───────────────────────

class TestAuthHealth:
    def test_login_success(self, admin_token):
        assert isinstance(admin_token, str) and len(admin_token) > 20

    def test_export_docx_requires_auth(self):
        r = requests.get(f"{BASE_URL}/api/events/export/whatever-id/docx", timeout=30)
        assert r.status_code in (401, 403), f"expected 401/403 got {r.status_code}"


# ─────────────────────── core flow ───────────────────────

class TestFasahDocxFlow:
    def test_export_docx_success_and_mime(self, admin_headers, analysis_id):
        aid = analysis_id["analysis_id"]
        r = requests.get(f"{BASE_URL}/api/events/export/{aid}/docx",
                         headers=admin_headers, timeout=120)
        assert r.status_code == 200, f"docx export failed: {r.status_code} {r.text[:300]}"
        ct = r.headers.get("Content-Type", "")
        assert DOCX_MIME in ct, f"unexpected content-type: {ct}"
        cd = r.headers.get("Content-Disposition", "")
        assert "attachment" in cd and ".docx" in cd
        assert "Weekly_Report_" in cd
        # Confirm it's a valid OOXML zip
        assert zipfile.is_zipfile(io.BytesIO(r.content))
        assert len(r.content) > 5000  # reasonable file size

    def test_docx_structure_headings_and_tables(self, admin_headers, analysis_id):
        aid = analysis_id["analysis_id"]
        r = requests.get(f"{BASE_URL}/api/events/export/{aid}/docx",
                         headers=admin_headers, timeout=120)
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.content))

        all_text = "\n".join(p.text for p in doc.paragraphs)
        for required in ("Notes on Last Week's Alerts",
                         "Concise Detailed Summary",
                         "Alerts Graph",
                         "Alert Summary"):
            assert required in all_text, f"Missing required heading: '{required}'"

        # At least 2 tables: banner (1x1) + alert summary (Nx4)
        assert len(doc.tables) >= 2, f"expected >=2 tables, got {len(doc.tables)}"
        # Locate the alert-summary table (4 cols, header text matches)
        summary_tbl = None
        for tbl in doc.tables:
            if len(tbl.columns) == 4 and len(tbl.rows) >= 1:
                hdr = [c.text.strip() for c in tbl.rows[0].cells]
                if "Health-Rules Name" in hdr[0]:
                    summary_tbl = tbl
                    break
        assert summary_tbl is not None, "Alert Summary 4-column table not found"
        hdr = [c.text.strip() for c in summary_tbl.rows[0].cells]
        assert hdr[0] == "Health-Rules Name"
        assert hdr[1] == "Critical Alert/Warning"
        assert hdr[2] == "Total"
        assert "Summary of Alert Triggered" in hdr[3]
        # Should have at least 1 data row
        assert len(summary_tbl.rows) >= 2, "no data rows in Alert Summary table"

    def test_branding_company_and_period(self, admin_headers, analysis_id):
        aid = analysis_id["analysis_id"]
        params = {"company": "Fasah", "period": "12 Apr - 18 Apr 2026"}
        r = requests.get(f"{BASE_URL}/api/events/export/{aid}/docx",
                         headers=admin_headers, params=params, timeout=120)
        assert r.status_code == 200
        doc = Document(io.BytesIO(r.content))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Banner is in a table cell so also check tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    all_text += "\n" + cell.text
        assert "Fasah" in all_text, "company name 'Fasah' not present in DOCX"
        assert "12 Apr - 18 Apr 2026" in all_text, "period string not present in DOCX"
        # Filename includes Fasah
        cd = r.headers.get("Content-Disposition", "")
        assert "Fasah" in cd

    def test_severity_color_coding_in_xml(self, admin_headers, analysis_id):
        """Verify <w:shd w:fill=...> differs between Critical and Warning rows."""
        aid = analysis_id["analysis_id"]
        r = requests.get(f"{BASE_URL}/api/events/export/{aid}/docx",
                         headers=admin_headers, timeout=120)
        assert r.status_code == 200
        with zipfile.ZipFile(io.BytesIO(r.content)) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        # Expected fills from fasah_report_service: critical=FCEAE8, warning=FDF4E3, info=EAF2F8
        # Header bg = 1F2937
        fills = set(re.findall(r'w:fill="([0-9A-Fa-f]{6})"', xml))
        assert "1F2937" in fills, "header navy fill missing"
        # At least one of the severity-specific fills must be present (since alert_rows non-empty)
        sev_fills = {"FCEAE8", "FDF4E3", "EAF2F8"}
        assert fills & sev_fills, f"no severity fill present, fills={fills}"
        # When test data has both Critical and Warning, both fills should appear
        # Data has 'critical' (Database timeout, CPU) and 'warning' (latency, memory)
        assert "FCEAE8" in fills, f"critical fill missing, fills={fills}"
        assert "FDF4E3" in fills, f"warning fill missing, fills={fills}"

    def test_404_for_unknown_analysis(self, admin_headers):
        r = requests.get(f"{BASE_URL}/api/events/export/non-existent-id-zzz/docx",
                         headers=admin_headers, timeout=30)
        assert r.status_code == 404
        assert "not found" in r.text.lower()


# ─────────────────────── existing flows still work ───────────────────────

class TestExistingExports:
    def test_excel_export_still_works(self, admin_headers, analysis_id):
        aid = analysis_id["analysis_id"]
        r = requests.get(f"{BASE_URL}/api/events/export/{aid}/excel",
                         headers=admin_headers, timeout=60)
        assert r.status_code == 200
        assert "spreadsheetml.sheet" in r.headers.get("Content-Type", "")
        assert zipfile.is_zipfile(io.BytesIO(r.content))

    def test_pdf_export_still_works(self, admin_headers, analysis_id):
        aid = analysis_id["analysis_id"]
        r = requests.get(f"{BASE_URL}/api/events/export/{aid}/pdf",
                         headers=admin_headers, timeout=60)
        assert r.status_code == 200
        assert r.headers.get("Content-Type", "").startswith("application/pdf")
        assert r.content[:4] == b"%PDF"


# ─────────────────────── robust parsing ───────────────────────

class TestRobustParsing:
    def test_mixed_case_headers(self, admin_headers):
        """Excel with mixed-case headers like 'TimeStamp', 'Service Name', 'Alert Description', 'Priority'."""
        rows = _sample_rows()  # the same row data
        xlsx = _make_xlsx(rows, headers=("TimeStamp", "Service Name", "Alert Description", "Priority"))
        files = {"file": ("TEST_mixedcase.xlsx", xlsx, "application/octet-stream")}
        r = requests.post(f"{BASE_URL}/api/events/upload",
                          headers=admin_headers, files=files, timeout=60)
        assert r.status_code == 200, f"mixed-case upload failed: {r.text}"
        body = r.json()
        assert body.get("total_events", 0) > 0, "events not parsed from mixed-case headers"
        upload_id = body["upload_id"]

        ar = requests.post(f"{BASE_URL}/api/events/analyze/{upload_id}",
                           headers=admin_headers, timeout=180)
        assert ar.status_code == 200, f"analyze failed: {ar.text}"
        aid = ar.json().get("analysis_id")
        assert aid
        # DOCX should still generate
        dr = requests.get(f"{BASE_URL}/api/events/export/{aid}/docx",
                          headers=admin_headers, timeout=120)
        assert dr.status_code == 200
        assert zipfile.is_zipfile(io.BytesIO(dr.content))


# ─────────────────────── empty / edge-case alerts ───────────────────────

class TestEmptyAlerts:
    def test_docx_generates_when_alerts_blank(self, admin_headers):
        """Empty/blank alert names → patterns.alert_frequency should be empty,
        DOCX must still 200 with the 'No alerts recorded' fallback."""
        rows = []
        base = datetime(2026, 1, 1, 9, 0, 0)
        for i in range(3):
            rows.append([(base + timedelta(minutes=i)).strftime("%Y-%m-%d %H:%M:%S"),
                         "svc", "", "info"])
        xlsx = _make_xlsx(rows)
        files = {"file": ("TEST_empty_alerts.xlsx", xlsx, "application/octet-stream")}
        r = requests.post(f"{BASE_URL}/api/events/upload",
                          headers=admin_headers, files=files, timeout=60)
        assert r.status_code == 200
        upload_id = r.json()["upload_id"]
        ar = requests.post(f"{BASE_URL}/api/events/analyze/{upload_id}",
                           headers=admin_headers, timeout=180)
        assert ar.status_code == 200, ar.text
        aid = ar.json().get("analysis_id")
        if not aid:
            pytest.skip("analyze did not return analysis_id for empty alerts")

        dr = requests.get(f"{BASE_URL}/api/events/export/{aid}/docx",
                          headers=admin_headers, timeout=120)
        assert dr.status_code == 200, f"docx failed for empty alerts: {dr.text[:300]}"
        doc = Document(io.BytesIO(dr.content))
        text = "\n".join(p.text for p in doc.paragraphs)
        # Must not crash; should have fallback message OR empty-table handling
        # The service writes "No alerts recorded for this period." when alert_rows empty.
        # If alert_freq has entries even for blanks (depends on parser), this is permissive.
        assert "Alert Summary" in text


# ─────────────────────── cleanup ───────────────────────

class TestZZZCleanup:
    def test_cleanup_test_uploads(self, admin_headers):
        """Best-effort cleanup of TEST_*.xlsx uploads."""
        try:
            r = requests.get(f"{BASE_URL}/api/events/uploads?limit=100",
                             headers=admin_headers, timeout=30)
            if r.status_code != 200:
                pytest.skip(f"cannot list uploads: {r.status_code}")
            for up in r.json().get("uploads", []):
                if str(up.get("filename", "")).startswith("TEST_"):
                    requests.delete(f"{BASE_URL}/api/events/upload/{up['id']}",
                                    headers=admin_headers, timeout=30)
        except Exception as e:
            pytest.skip(f"cleanup error: {e}")

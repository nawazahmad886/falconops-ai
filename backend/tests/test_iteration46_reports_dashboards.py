"""
Iteration 46 - AI Weekly Report Generator + Custom Dashboard Builder
Tests /api/weekly-reports/* and /api/custom-dashboards/* endpoints
"""
import io
import os
import pytest
import requests

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "https://health-rules-engine.preview.emergentagent.com").rstrip("/")
ADMIN_EMAIL = "admin@falconapps.com"
ADMIN_PASS = "Admin@123"


# ---------- Fixtures ----------
@pytest.fixture(scope="module")
def token():
    r = requests.post(f"{BASE_URL}/api/auth/login",
                      json={"email": ADMIN_EMAIL, "password": ADMIN_PASS}, timeout=30)
    assert r.status_code == 200, f"Login failed: {r.status_code} {r.text}"
    return r.json()["access_token"]


@pytest.fixture(scope="module")
def headers(token):
    return {"Authorization": f"Bearer {token}"}


# ---------- Weekly Reports ----------
class TestWeeklyReports:
    _report_id = None

    def test_auto_generate(self, headers):
        r = requests.post(f"{BASE_URL}/api/weekly-reports/generate/auto",
                          json={"days": 7, "period": ""}, headers=headers, timeout=120)
        assert r.status_code == 200, r.text
        data = r.json()
        assert "report_id" in data
        assert "ai_summary" in data
        assert isinstance(data.get("alerts"), list)
        assert isinstance(data.get("total_alerts"), int)
        TestWeeklyReports._report_id = data["report_id"]

    def test_json_generate(self, headers):
        payload = {
            "period": "test-week",
            "alerts": [
                {"rule_name": "TEST_CPU_High", "severity": "critical", "count": 5, "summary": "CPU high"},
                {"rule_name": "TEST_Mem_Warn", "severity": "warning", "count": 2, "summary": "Mem warn"},
            ],
        }
        r = requests.post(f"{BASE_URL}/api/weekly-reports/generate/json",
                          json=payload, headers=headers, timeout=60)
        assert r.status_code == 200, r.text
        data = r.json()
        assert "report_id" in data
        assert data["total_alerts"] == 2
        assert data["critical_count"] >= 1

    def test_upload_docx(self, headers):
        # Build a minimal DOCX with python-docx
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc = Document()
        doc.add_heading("Weekly Alerts", 0)
        table = doc.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = "Rule Name"
        table.rows[0].cells[1].text = "Severity"
        table.rows[0].cells[2].text = "Count"
        table.rows[0].cells[3].text = "Summary"
        table.rows[1].cells[0].text = "TEST_UPLOAD_RULE"
        table.rows[1].cells[1].text = "critical"
        table.rows[1].cells[2].text = "3"
        table.rows[1].cells[3].text = "Upload test"
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        files = {"file": ("test.docx", buf.getvalue(),
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{BASE_URL}/api/weekly-reports/generate/upload",
                          files=files, headers=headers, timeout=90)
        assert r.status_code == 200, r.text
        data = r.json()
        assert "report_id" in data
        assert "ai_summary" in data

    def test_list_reports(self, headers):
        r = requests.get(f"{BASE_URL}/api/weekly-reports/list", headers=headers, timeout=30)
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data, list)
        assert len(data) >= 1

    def test_get_single_report(self, headers):
        assert TestWeeklyReports._report_id, "No report_id from auto test"
        r = requests.get(f"{BASE_URL}/api/weekly-reports/{TestWeeklyReports._report_id}",
                         headers=headers, timeout=30)
        assert r.status_code == 200
        data = r.json()
        assert data.get("report_id") == TestWeeklyReports._report_id

    def test_download_docx(self, headers):
        assert TestWeeklyReports._report_id
        r = requests.get(f"{BASE_URL}/api/weekly-reports/{TestWeeklyReports._report_id}/download/docx",
                         headers=headers, timeout=60)
        assert r.status_code == 200
        ct = r.headers.get("content-type", "")
        assert "wordprocessingml" in ct or "octet-stream" in ct, f"content-type={ct}"
        assert len(r.content) > 500

    def test_download_excel(self, headers):
        assert TestWeeklyReports._report_id
        r = requests.get(f"{BASE_URL}/api/weekly-reports/{TestWeeklyReports._report_id}/download/excel",
                         headers=headers, timeout=60)
        assert r.status_code == 200
        ct = r.headers.get("content-type", "")
        assert "spreadsheetml" in ct or "octet-stream" in ct, f"content-type={ct}"
        assert len(r.content) > 500


# ---------- Custom Dashboards ----------
class TestCustomDashboards:
    _dash_id = None

    def test_widget_catalog(self, headers):
        r = requests.get(f"{BASE_URL}/api/custom-dashboards/widgets/catalog",
                         headers=headers, timeout=30)
        assert r.status_code == 200
        data = r.json()
        assert "widgets" in data
        assert len(data["widgets"]) == 10
        types = {w["type"] for w in data["widgets"]}
        expected = {"soc_feed", "uptime", "threats", "billing", "sla",
                    "ai_agents", "alerts", "incidents", "metrics", "tenants"}
        assert expected.issubset(types)

    @pytest.mark.parametrize("wtype", [
        "soc_feed", "uptime", "threats", "billing", "sla",
        "ai_agents", "alerts", "incidents", "metrics", "tenants",
    ])
    def test_widget_data(self, headers, wtype):
        r = requests.get(f"{BASE_URL}/api/custom-dashboards/widgets/data/{wtype}",
                         headers=headers, timeout=30)
        assert r.status_code == 200, f"{wtype}: {r.text}"
        data = r.json()
        assert isinstance(data, dict)

    def test_create_dashboard(self, headers):
        payload = {
            "name": "TEST_Dashboard",
            "description": "iter46 test",
            "widgets": [
                {"i": "w1", "x": 0, "y": 0, "w": 4, "h": 3, "widget_type": "uptime", "title": "Uptime"},
                {"i": "w2", "x": 4, "y": 0, "w": 4, "h": 3, "widget_type": "metrics", "title": "Metrics"},
            ],
        }
        r = requests.post(f"{BASE_URL}/api/custom-dashboards/create",
                          json=payload, headers=headers, timeout=30)
        assert r.status_code == 200, r.text
        data = r.json()
        assert "dashboard_id" in data
        assert data["name"] == "TEST_Dashboard"
        assert len(data["widgets"]) == 2
        TestCustomDashboards._dash_id = data["dashboard_id"]

    def test_list_dashboards(self, headers):
        r = requests.get(f"{BASE_URL}/api/custom-dashboards/list", headers=headers, timeout=30)
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data, list)
        ids = [d.get("dashboard_id") for d in data]
        assert TestCustomDashboards._dash_id in ids

    def test_get_dashboard(self, headers):
        r = requests.get(f"{BASE_URL}/api/custom-dashboards/{TestCustomDashboards._dash_id}",
                         headers=headers, timeout=30)
        assert r.status_code == 200
        assert r.json()["dashboard_id"] == TestCustomDashboards._dash_id

    def test_update_dashboard(self, headers):
        payload = {
            "name": "TEST_Dashboard_Updated",
            "description": "updated",
            "widgets": [
                {"i": "w1", "x": 0, "y": 0, "w": 6, "h": 4, "widget_type": "alerts", "title": "Alerts"},
            ],
        }
        r = requests.put(f"{BASE_URL}/api/custom-dashboards/{TestCustomDashboards._dash_id}",
                         json=payload, headers=headers, timeout=30)
        assert r.status_code == 200
        data = r.json()
        assert data["name"] == "TEST_Dashboard_Updated"
        assert len(data["widgets"]) == 1

        # GET verify
        g = requests.get(f"{BASE_URL}/api/custom-dashboards/{TestCustomDashboards._dash_id}",
                        headers=headers, timeout=30)
        assert g.status_code == 200
        assert g.json()["name"] == "TEST_Dashboard_Updated"

    def test_delete_dashboard(self, headers):
        r = requests.delete(f"{BASE_URL}/api/custom-dashboards/{TestCustomDashboards._dash_id}",
                            headers=headers, timeout=30)
        assert r.status_code == 200
        # GET -> 404
        g = requests.get(f"{BASE_URL}/api/custom-dashboards/{TestCustomDashboards._dash_id}",
                        headers=headers, timeout=30)
        assert g.status_code == 404

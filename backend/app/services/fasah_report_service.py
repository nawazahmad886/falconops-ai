"""
FalconOps AI — Fasah Weekly Report DOCX Generator
Generates a professional weekly report in the exact Fasah template format:
- Title block with bilingual headings + date range
- Notes on last week's alerts (executive bullets)
- Concise detailed summary per point
- Alert frequency bar chart
- Severity pie chart
- Main alert summary table (Health-Rules | Severity | Total | Summary) with colour coding

Maps event-analysis JSON (from EventAnalyzer.ai_analyze) → docx bytes.
"""
import io
import logging
from datetime import datetime, timezone, timedelta
from typing import Dict, List, Optional

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Fasah colour palette (from original template)
COLOR_CRITICAL = RGBColor(0xC0, 0x39, 0x2B)   # deep red
COLOR_WARNING  = RGBColor(0xE6, 0x8A, 0x19)   # amber
COLOR_INFO     = RGBColor(0x27, 0x74, 0xA6)   # blue
COLOR_HEADER_BG = "1F2937"                    # navy header bg (hex)
COLOR_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_SUBTLE   = RGBColor(0x60, 0x60, 0x60)


# ──────────────────────────────────────────────
#  Helpers
# ──────────────────────────────────────────────

def _set_cell_background(cell, hex_color: str):
    """Apply a solid fill colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_page_number_footer(doc: Document):
    """Add 'Page X of Y' centered footer to every section."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_SUBTLE

        # PAGE field
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_begin, instr, fld_end])
        run.add_text(" / ")
        fld_begin2 = OxmlElement("w:fldChar"); fld_begin2.set(qn("w:fldCharType"), "begin")
        instr2 = OxmlElement("w:instrText"); instr2.text = "NUMPAGES"
        fld_end2 = OxmlElement("w:fldChar"); fld_end2.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_begin2, instr2, fld_end2])


def _severity_color(severity: str) -> RGBColor:
    s = (severity or "").lower()
    if "crit" in s:
        return COLOR_CRITICAL
    if "warn" in s or "high" in s or "major" in s:
        return COLOR_WARNING
    return COLOR_INFO


def _severity_bg(severity: str) -> str:
    s = (severity or "").lower()
    if "crit" in s:
        return "FCEAE8"
    if "warn" in s or "high" in s or "major" in s:
        return "FDF4E3"
    return "EAF2F8"


def _resolve_period(period: Optional[str]) -> str:
    if period:
        return period
    now = datetime.now(timezone.utc)
    week_ago = now - timedelta(days=7)
    return f"{week_ago.strftime('%d %b')} – {now.strftime('%d %b %Y')}"


# ──────────────────────────────────────────────
#  Charts
# ──────────────────────────────────────────────

def _make_alert_bar_chart(alert_freq: List[Dict], out: io.BytesIO):
    if not alert_freq:
        return None
    top = alert_freq[:10]
    labels = [a.get("alert", "?")[:40] for a in top]
    counts = [a.get("count", 0) for a in top]
    fig, ax = plt.subplots(figsize=(7.4, 3.6), dpi=160)
    ax.barh(labels[::-1], counts[::-1], color="#C0392B", edgecolor="#8E2A1E")
    ax.set_xlabel("Occurrences", fontsize=9)
    ax.set_title("Top 10 Alert Types (Last 7 Days)", fontsize=11, fontweight="bold", pad=10)
    ax.tick_params(axis="y", labelsize=8)
    ax.tick_params(axis="x", labelsize=8)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.grid(axis="x", alpha=0.15)
    plt.tight_layout()
    fig.savefig(out, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    out.seek(0)
    return out


def _make_severity_pie(sev_dist: Dict[str, int], out: io.BytesIO):
    if not sev_dist:
        return None
    labels, values, colors = [], [], []
    palette = {"critical": "#C0392B", "warning": "#E68A19", "info": "#2774A6"}
    for k in ("critical", "warning", "info"):
        if sev_dist.get(k):
            labels.append(f"{k.title()} ({sev_dist[k]})")
            values.append(sev_dist[k])
            colors.append(palette[k])
    if not values:
        return None
    fig, ax = plt.subplots(figsize=(4.2, 3.6), dpi=160)
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, colors=colors, autopct="%1.0f%%",
        startangle=90, wedgeprops=dict(edgecolor="white", linewidth=2),
        textprops=dict(fontsize=9),
    )
    for t in autotexts:
        t.set_color("white"); t.set_fontweight("bold")
    ax.set_title("Severity Distribution", fontsize=11, fontweight="bold", pad=10)
    plt.tight_layout()
    fig.savefig(out, format="png", dpi=160, bbox_inches="tight")
    plt.close(fig)
    out.seek(0)
    return out


# ──────────────────────────────────────────────
#  Data mapping (event-analysis → Fasah sections)
# ──────────────────────────────────────────────

def _extract_alert_rows(analysis: Dict, events: List[Dict], max_rows: int = 25) -> List[Dict]:
    """Aggregate alert_frequency (name, count) and enrich each row with inferred
    severity + a concise summary, using events as source of truth."""
    patterns = analysis.get("patterns", {}) or {}
    alert_freq = patterns.get("alert_frequency", []) or []

    # Build an alert → (severity, sample events) index
    by_alert: Dict[str, Dict] = {}
    for ev in events:
        alert_name = str(ev.get("alert") or ev.get("alert_description") or ev.get("message") or ev.get("description") or "").strip()
        if not alert_name:
            continue
        alert_name_key = alert_name[:100]
        entry = by_alert.setdefault(alert_name_key, {"severities": [], "services": [], "samples": []})
        sev = str(ev.get("severity") or ev.get("priority") or "info").strip().lower()
        entry["severities"].append(sev)
        svc = str(ev.get("service") or ev.get("application") or "").strip()
        if svc:
            entry["services"].append(svc)
        if len(entry["samples"]) < 3:
            entry["samples"].append(ev)

    # Build Fasah-style rows from alert_freq (which is ordered by count desc)
    rows: List[Dict] = []
    for af in alert_freq[:max_rows]:
        name = af.get("alert") or "Unknown Alert"
        count = af.get("count", 0)
        meta = by_alert.get(name, {})
        severities = meta.get("severities", ["info"])
        # Dominant severity
        if any("crit" in s for s in severities):
            sev_label = "Critical"
        elif any(("warn" in s or "high" in s or "major" in s) for s in severities):
            sev_label = "Warning"
        else:
            sev_label = "Info"

        # Summary text
        services = list(dict.fromkeys(meta.get("services", [])))[:3]
        svc_blurb = f"Affected services: {', '.join(services)}. " if services else ""
        summary = f"{svc_blurb}Triggered {count} time(s) in the past week."

        rows.append({
            "rule_name": name,
            "severity": sev_label,
            "count": count,
            "summary": summary,
        })
    return rows


def _build_executive_bullets(analysis: Dict, alert_rows: List[Dict]) -> List[str]:
    """Build the Fasah-style executive observation bullets for 'Notes on Last Week's Alerts'."""
    patterns = analysis.get("patterns", {}) or {}
    summary = analysis.get("summary", {}) or {}
    bullets: List[str] = []

    total = patterns.get("total_events") or summary.get("total_events") or 0
    unique = patterns.get("unique_alerts", 0)
    critical_count = sum(1 for r in alert_rows if r["severity"] == "Critical")
    warning_count = sum(1 for r in alert_rows if r["severity"] == "Warning")

    bullets.append(
        f"Monitored platform recorded {total} alert event(s) across {unique} unique alert types during the period."
    )
    if critical_count:
        bullets.append(f"{critical_count} distinct critical-severity alert pattern(s) require immediate attention.")
    if warning_count:
        bullets.append(f"{warning_count} warning-severity pattern(s) should be reviewed and tuned to reduce noise.")

    top_services = patterns.get("service_frequency", [])[:3]
    if top_services:
        svc_list = ", ".join(f"{s.get('service')} ({s.get('count')})" for s in top_services)
        bullets.append(f"Most affected services: {svc_list}.")

    top_alerts = alert_rows[:3]
    if top_alerts:
        alert_list = "; ".join(f"‘{a['rule_name'][:50]}’ × {a['count']}" for a in top_alerts)
        bullets.append(f"Top recurring alerts driving the week's volume: {alert_list}.")

    return bullets


# ──────────────────────────────────────────────
#  Main generator
# ──────────────────────────────────────────────

def generate_fasah_weekly_report(
    analysis: Dict,
    events: List[Dict],
    period: Optional[str] = None,
    company_name: str = "Fasah",
    report_title: str = "Weekly AIOps Report",
) -> bytes:
    """Return DOCX bytes in Fasah Weekly Report format."""
    doc = Document()

    # Page margins (portrait, narrow)
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    _add_page_number_footer(doc)

    # ─── Header banner ───
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = True
    cell = banner.rows[0].cells[0]
    _set_cell_background(cell, COLOR_HEADER_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name} — Platform Alerts  |  تنبيهات المنصة")
    run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = COLOR_HEADER_TEXT
    sub = cell.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Report Date · {datetime.now(timezone.utc).strftime('%d/%m/%Y')}    ·    Period · {_resolve_period(period)}")
    sub_run.font.size = Pt(10); sub_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

    doc.add_paragraph()

    # ─── Title ───
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(report_title)
    title_run.font.size = Pt(18); title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_paragraph()

    # ─── Section 1: Executive notes (Notes on Last Week's Alerts) ───
    h = doc.add_heading("Notes on Last Week's Alerts", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    alert_rows = _extract_alert_rows(analysis, events, max_rows=30)
    bullets = _build_executive_bullets(analysis, alert_rows)

    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b).font.size = Pt(11)

    # ─── Section 2: Concise detailed summary (AI summary) ───
    doc.add_paragraph()
    h2 = doc.add_heading("Concise Detailed Summary", level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    ai_text = analysis.get("ai_analysis") or analysis.get("executive_summary") or ""
    if not ai_text:
        # Build a fallback executive summary
        patterns = analysis.get("patterns", {})
        total = patterns.get("total_events", 0)
        sev = patterns.get("severity_distribution", {})
        ai_text = (
            f"During the reporting period the platform recorded {total} alert event(s). "
            f"Severity mix — Critical: {sev.get('critical', 0)}, Warning: {sev.get('warning', 0)}, "
            f"Info: {sev.get('info', 0)}. "
            "Detailed per-rule analysis follows in the Alert Summary table below."
        )
    p = doc.add_paragraph(ai_text.strip())
    for run in p.runs:
        run.font.size = Pt(11)

    # ─── Section 3: Alerts graph ───
    doc.add_paragraph()
    h3 = doc.add_heading("Alerts Graph", level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    patterns = analysis.get("patterns", {}) or {}
    try:
        bar_buf = io.BytesIO()
        if _make_alert_bar_chart(patterns.get("alert_frequency", []), bar_buf):
            doc.add_picture(bar_buf, width=Inches(6.8))
    except Exception as e:
        logger.warning("Failed to render bar chart: %s", e)

    try:
        pie_buf = io.BytesIO()
        if _make_severity_pie(patterns.get("severity_distribution", {}), pie_buf):
            doc.add_picture(pie_buf, width=Inches(4.2))
    except Exception as e:
        logger.warning("Failed to render pie chart: %s", e)

    # ─── Section 4: Alert Summary table ───
    doc.add_paragraph()
    h4 = doc.add_heading("Alert Summary", level=1)
    for run in h4.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    if not alert_rows:
        doc.add_paragraph("No alerts recorded for this period.")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Column widths
        widths = [Inches(2.2), Inches(1.3), Inches(0.7), Inches(2.6)]
        for i, w in enumerate(widths):
            for c in table.columns[i].cells:
                c.width = w

        # Header row
        hdr_cells = table.rows[0].cells
        headers = ["Health-Rules Name", "Critical Alert/Warning", "Total", "Summary of Alert Triggered and Action/Issues"]
        for i, text in enumerate(headers):
            _set_cell_background(hdr_cells[i], COLOR_HEADER_BG)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = hdr_cells[i].paragraphs[0].add_run(text)
            run.font.bold = True; run.font.size = Pt(10); run.font.color.rgb = COLOR_HEADER_TEXT

        # Data rows
        for row_data in alert_rows:
            row = table.add_row().cells
            row[0].text = row_data["rule_name"][:120]
            _set_cell_background(row[1], _severity_bg(row_data["severity"]))
            sev_run = row[1].paragraphs[0].add_run(row_data["severity"])
            sev_run.font.bold = True
            sev_run.font.color.rgb = _severity_color(row_data["severity"])
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].text = str(row_data["count"])
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[3].text = row_data["summary"]
            # Apply font size to all cells
            for cell in row:
                for para in cell.paragraphs:
                    for r in para.runs:
                        if not r.font.size:
                            r.font.size = Pt(9)

    # ─── Section 5: AI suggestions (if present) ───
    suggestions = analysis.get("suggestions") or []
    if suggestions:
        doc.add_paragraph()
        h5 = doc.add_heading("Recommended Actions", level=1)
        for run in h5.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        for s in suggestions[:8]:
            action = s.get("action") or s.get("title") or str(s)[:200]
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(action)
            run.font.size = Pt(11)
            desc = s.get("description") or s.get("details")
            if desc:
                sub = doc.add_paragraph()
                sub.paragraph_format.left_indent = Cm(1.2)
                sub_run = sub.add_run(desc[:400])
                sub_run.font.size = Pt(9); sub_run.font.color.rgb = COLOR_SUBTLE

    # ─── Footer line ───
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = p.add_run(f"Generated by FalconOps AI · {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    fr.font.size = Pt(9); fr.font.color.rgb = COLOR_SUBTLE

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

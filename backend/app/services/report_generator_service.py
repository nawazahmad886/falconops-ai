"""
FalconOps AI - Weekly Report Generator Service
Parses DOCX/CSV/JSON, runs AI analysis, generates DOCX + Excel + PDF exports
Matches the Fasah Weekly Report format + Enterprise (Datadog/Splunk style) PDF
"""
import uuid
import io
import os
import logging
import tempfile
from datetime import datetime, timezone, timedelta
from typing import Dict, List, Optional

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, Reference

from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

from ..core.database import db

logger = logging.getLogger(__name__)

UPLOAD_DIR = "/tmp/falconops_reports"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ======================== DOCX PARSER ========================

def parse_docx(content: bytes) -> Dict:
    """Parse uploaded DOCX report and extract alert data"""
    doc = Document(io.BytesIO(content))
    alerts = []
    summary_lines = []
    in_table = False

    # Extract paragraphs for summary
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 10:
            summary_lines.append(text)

    # Extract tables
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 3:
                alert = {
                    "rule_name": cells[0] if len(cells) > 0 else "",
                    "severity": "critical" if "critical" in (cells[1] if len(cells) > 1 else "").lower() else "warning",
                    "count": 0,
                    "summary": cells[3] if len(cells) > 3 else cells[-1],
                }
                try:
                    alert["count"] = int(cells[2]) if len(cells) > 2 and cells[2].isdigit() else 1
                except (ValueError, IndexError):
                    alert["count"] = 1
                if alert["rule_name"]:
                    alerts.append(alert)

    return {
        "source": "docx_upload",
        "alerts": alerts,
        "total_alerts": len(alerts),
        "critical_count": sum(1 for a in alerts if a["severity"] == "critical"),
        "warning_count": sum(1 for a in alerts if a["severity"] == "warning"),
        "total_occurrences": sum(a["count"] for a in alerts),
        "summary_text": "\n".join(summary_lines[:20]),
    }


def parse_csv_data(rows: List[Dict]) -> Dict:
    """Parse CSV/JSON array of alerts"""
    alerts = []
    for row in rows:
        alerts.append({
            "rule_name": row.get("rule_name", row.get("name", row.get("alert", ""))),
            "severity": row.get("severity", "warning"),
            "count": int(row.get("count", row.get("total", 1))),
            "summary": row.get("summary", row.get("message", row.get("description", ""))),
        })
    return {
        "source": "json_upload",
        "alerts": alerts,
        "total_alerts": len(alerts),
        "critical_count": sum(1 for a in alerts if a["severity"] == "critical"),
        "warning_count": sum(1 for a in alerts if a["severity"] == "warning"),
        "total_occurrences": sum(a["count"] for a in alerts),
    }


def _normalize_header(h: str) -> str:
    return (h or "").strip().lower().replace(" ", "_")


def _row_to_alert(row: Dict) -> Dict:
    """Flexible mapping from arbitrary column names → alert schema."""
    # Lowercase all keys
    r = {_normalize_header(k): v for k, v in row.items() if v is not None}

    # Rule name
    rule_name = (
        r.get("rule_name") or r.get("rule") or r.get("name") or r.get("alert") or
        r.get("alert_name") or r.get("alertname") or r.get("title") or ""
    )

    # Severity
    sev_raw = str(r.get("severity") or r.get("priority") or r.get("level") or "warning").lower().strip()
    if sev_raw in ("p1", "p0", "sev1", "sev0", "crit"):
        severity = "critical"
    elif sev_raw in ("p2", "sev2", "err", "error"):
        severity = "high"
    elif sev_raw in ("p3", "sev3", "warn"):
        severity = "warning"
    elif sev_raw in ("p4", "sev4", "informational"):
        severity = "info"
    elif sev_raw in ("critical", "high", "warning", "info"):
        severity = sev_raw
    else:
        severity = "warning"

    # Count / occurrences
    count_val = r.get("count") or r.get("occurrences") or r.get("total") or r.get("hits") or 1
    try:
        count = int(float(count_val))
    except (ValueError, TypeError):
        count = 1

    # Summary
    summary = (
        r.get("summary") or r.get("message") or r.get("description") or
        r.get("details") or r.get("notes") or ""
    )

    return {
        "rule_name": str(rule_name).strip()[:200],
        "severity": severity,
        "count": max(1, count),
        "summary": str(summary).strip()[:500],
    }


def parse_excel(content: bytes) -> Dict:
    """Parse .xlsx file — first sheet, first row is header. Rows must include at least rule_name."""
    import io as _io
    from openpyxl import load_workbook

    wb = load_workbook(filename=_io.BytesIO(content), data_only=True, read_only=True)
    ws = wb.active
    rows_iter = ws.iter_rows(values_only=True)
    try:
        headers_raw = next(rows_iter)
    except StopIteration:
        return {"source": "excel_upload", "alerts": [], "total_alerts": 0,
                "critical_count": 0, "warning_count": 0, "total_occurrences": 0}

    headers = [_normalize_header(str(h) if h is not None else f"col{i}") for i, h in enumerate(headers_raw)]
    alerts = []
    for row in rows_iter:
        if row is None or all(c is None or c == "" for c in row):
            continue
        row_dict = dict(zip(headers, row))
        alert = _row_to_alert(row_dict)
        if alert["rule_name"]:
            alerts.append(alert)

    return {
        "source": "excel_upload",
        "alerts": alerts,
        "total_alerts": len(alerts),
        "critical_count": sum(1 for a in alerts if a["severity"] == "critical"),
        "warning_count": sum(1 for a in alerts if a["severity"] == "warning"),
        "total_occurrences": sum(a["count"] for a in alerts),
    }


def parse_csv_bytes(content: bytes) -> Dict:
    """Parse .csv file bytes."""
    import csv as _csv
    import io as _io

    text = content.decode("utf-8-sig", errors="replace")
    reader = _csv.DictReader(_io.StringIO(text))
    alerts = []
    for row in reader:
        alert = _row_to_alert(row)
        if alert["rule_name"]:
            alerts.append(alert)

    return {
        "source": "csv_upload",
        "alerts": alerts,
        "total_alerts": len(alerts),
        "critical_count": sum(1 for a in alerts if a["severity"] == "critical"),
        "warning_count": sum(1 for a in alerts if a["severity"] == "warning"),
        "total_occurrences": sum(a["count"] for a in alerts),
    }


# ======================== AUTO-FETCH FROM SOC ========================

async def fetch_from_soc(days: int = 7) -> Dict:
    """Auto-fetch alerts from the SOC engine for the past N days"""
    cutoff = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()

    # Get uptime alerts
    uptime_alerts = await db.uptime_alerts.find(
        {"timestamp": {"$gte": cutoff}}, {"_id": 0}
    ).to_list(500)

    # Get security threats
    threats = await db.security_threats.find(
        {"timestamp": {"$gte": cutoff}}, {"_id": 0}
    ).to_list(500)

    # Get SOC events
    soc_events = await db.soc_events.find(
        {"timestamp": {"$gte": cutoff}}, {"_id": 0}
    ).to_list(500)

    # Group by rule/type
    rule_map = {}
    for a in uptime_alerts:
        key = f"Uptime: {a.get('monitor_name', 'Unknown')}"
        if key not in rule_map:
            rule_map[key] = {"rule_name": key, "severity": "critical" if a.get("alert_type") == "down" else "warning", "count": 0, "summary": f"Monitor {a.get('alert_type', '')} - {a.get('url', '')}"}
        rule_map[key]["count"] += 1

    for t in threats:
        key = t.get("event_type", t.get("message", "Security Threat"))[:60]
        if key not in rule_map:
            rule_map[key] = {"rule_name": key, "severity": t.get("severity", "warning"), "count": 0, "summary": t.get("message", "")}
        rule_map[key]["count"] += 1

    for e in soc_events:
        key = f"{e.get('source', 'soc')}: {e.get('service', 'Unknown')}"
        if key not in rule_map:
            rule_map[key] = {"rule_name": key, "severity": e.get("severity", "info"), "count": 0, "summary": e.get("message", "")}
        rule_map[key]["count"] += 1

    alerts = sorted(rule_map.values(), key=lambda x: (-{"critical": 3, "high": 2, "warning": 1}.get(x["severity"], 0), -x["count"]))

    return {
        "source": "soc_auto_fetch",
        "period_days": days,
        "alerts": alerts,
        "total_alerts": len(alerts),
        "critical_count": sum(1 for a in alerts if a["severity"] == "critical"),
        "warning_count": sum(1 for a in alerts if a["severity"] in ("warning", "high")),
        "total_occurrences": sum(a["count"] for a in alerts),
    }


# ======================== AI ANALYSIS ========================

async def fetch_sla_metrics(days: int = 7) -> Dict:
    """Pull live SLA / MTTR / uptime metrics from the platform"""
    cutoff = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
    try:
        uptime_monitors = await db.uptime_monitors.find({}, {"_id": 0}).to_list(500)
        incidents = await db.incidents.find({"created_at": {"$gte": cutoff}}, {"_id": 0}).to_list(500)
        threats = await db.security_threats.find({"timestamp": {"$gte": cutoff}}, {"_id": 0}).to_list(500)
        alerts = await db.alerts.find({"created_at": {"$gte": cutoff}}, {"_id": 0}).to_list(500)

        total = max(len(uptime_monitors), 1)
        up = sum(1 for m in uptime_monitors if m.get("status") == "up")
        avg_uptime = round(sum(m.get("uptime_pct", 100.0) for m in uptime_monitors) / total, 2) if total else 100.0

        # MTTR: average (resolved_at - created_at) in minutes
        mttr_minutes = 0.0
        resolved = [i for i in incidents if i.get("status") in ("resolved", "closed") and i.get("resolved_at")]
        if resolved:
            total_min = 0.0
            for i in resolved:
                try:
                    start = datetime.fromisoformat(i["created_at"].replace("Z", "+00:00"))
                    end = datetime.fromisoformat(i["resolved_at"].replace("Z", "+00:00"))
                    total_min += (end - start).total_seconds() / 60.0
                except Exception:
                    pass
            mttr_minutes = round(total_min / max(len(resolved), 1), 1)

        critical_threats = sum(1 for t in threats if t.get("severity") == "critical")
        risk = "Low"
        if critical_threats > 10 or avg_uptime < 99.0:
            risk = "High"
        elif critical_threats > 3 or avg_uptime < 99.5:
            risk = "Medium"

        return {
            "uptime_pct": avg_uptime,
            "monitors_total": total,
            "monitors_up": up,
            "monitors_down": total - up,
            "incidents_total": len(incidents),
            "incidents_resolved": len(resolved),
            "threats_total": len(threats),
            "threats_critical": critical_threats,
            "alerts_total": len(alerts),
            "mttr_minutes": mttr_minutes,
            "risk_posture": risk,
            "sla_compliance": "Compliant" if avg_uptime >= 99.9 else "At Risk" if avg_uptime >= 99.0 else "Breached",
        }
    except Exception as e:
        logger.error(f"SLA metrics fetch failed: {e}")
        return {
            "uptime_pct": 0, "monitors_total": 0, "monitors_up": 0, "monitors_down": 0,
            "incidents_total": 0, "incidents_resolved": 0, "threats_total": 0, "threats_critical": 0,
            "alerts_total": 0, "mttr_minutes": 0, "risk_posture": "Unknown", "sla_compliance": "Unknown",
        }


async def fetch_tenant_branding(tenant_id: Optional[str]) -> Dict:
    """Fetch tenant branding config (logo, colors, footer)"""
    default = {
        "company_name": "FalconOps AI",
        "logo_url": None,
        "primary_color": "#00E0FF",
        "secondary_color": "#F5B841",
        "footer_text": "Confidential - FalconOps AI SOC Report",
    }
    if not tenant_id:
        return default
    try:
        t = await db.tenants.find_one({"id": tenant_id}, {"_id": 0})
        if not t:
            return default
        branding = t.get("branding", {}) or {}
        return {
            "company_name": t.get("name") or default["company_name"],
            "logo_url": branding.get("logo_url"),
            "primary_color": branding.get("primary_color") or default["primary_color"],
            "secondary_color": branding.get("secondary_color") or default["secondary_color"],
            "footer_text": branding.get("footer_text") or f"Confidential - {t.get('name', 'ACME')} SOC",
        }
    except Exception as e:
        logger.error(f"Tenant branding fetch failed: {e}")
        return default


async def generate_ai_summary(parsed: Dict, sla: Optional[Dict] = None, executive: bool = False) -> str:
    """Use AI agents to generate report summary. executive=True uses CSO-level prompt."""
    try:
        from .ai_agents_service import run_agent
        data = {
            "report_type": "executive_cso_briefing" if executive else "weekly_alert_summary",
            "total_alerts": parsed["total_alerts"],
            "critical": parsed["critical_count"],
            "warnings": parsed["warning_count"],
            "total_occurrences": parsed.get("total_occurrences", 0),
            "top_alerts": [f"{a['rule_name']} ({a['severity']}, count:{a['count']})" for a in parsed["alerts"][:10]],
        }
        if sla:
            data["sla_metrics"] = sla
            data["instruction"] = (
                "You are a Chief Security Officer (CSO) drafting a weekly executive briefing. "
                "Deliver a concise, leadership-grade SOC report covering: "
                "1) Business Impact Summary (2-3 sentences). "
                "2) Key Incidents (top 5 numbered, each 1 line). "
                "3) Risk Posture (clearly marked Low / Medium / High with justification). "
                "4) SLA Compliance (with the numerical uptime and MTTR). "
                "5) Recommendations for Leadership (3-5 bullet action items). "
                "Use short paragraphs, no jargon. Total length under 500 words."
            )
        result = await run_agent("summarizer", data, use_memory=True)
        return result.get("analysis", "")
    except Exception as e:
        logger.error(f"AI summary failed: {e}")
        return "AI summary generation failed. Please review alerts manually."


# ======================== CHART GENERATOR ========================

def _chart_palette(primary: str = "#00E0FF"):
    return ["#DC2626", "#F97316", "#F59E0B", "#3B82F6", "#10B981", "#8B5CF6", "#EC4899"]


def generate_severity_chart(parsed: Dict) -> Optional[str]:
    """Bar chart of alerts by severity"""
    try:
        sev_counts = {}
        for a in parsed.get("alerts", []):
            s = a.get("severity", "info").capitalize()
            sev_counts[s] = sev_counts.get(s, 0) + a.get("count", 1)
        if not sev_counts:
            return None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png", dir=UPLOAD_DIR)
        tmp.close()
        plt.figure(figsize=(7, 4), facecolor="white")
        colors_map = {"Critical": "#DC2626", "High": "#F97316", "Warning": "#F59E0B", "Info": "#3B82F6"}
        bar_colors = [colors_map.get(k, "#64748B") for k in sev_counts.keys()]
        plt.bar(list(sev_counts.keys()), list(sev_counts.values()), color=bar_colors)
        plt.title("Alert Severity Distribution", fontsize=14, fontweight="bold")
        plt.ylabel("Occurrences")
        plt.tight_layout()
        plt.savefig(tmp.name, dpi=110, bbox_inches="tight")
        plt.close()
        return tmp.name
    except Exception as e:
        logger.error(f"Severity chart failed: {e}")
        return None


def generate_top_rules_chart(parsed: Dict) -> Optional[str]:
    """Horizontal bar chart of top-10 rules by count"""
    try:
        alerts = sorted(parsed.get("alerts", []), key=lambda x: -x.get("count", 0))[:10]
        if not alerts:
            return None
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png", dir=UPLOAD_DIR)
        tmp.close()
        names = [a["rule_name"][:35] for a in alerts]
        counts = [a.get("count", 0) for a in alerts]
        plt.figure(figsize=(8, 5), facecolor="white")
        plt.barh(names[::-1], counts[::-1], color="#00E0FF")
        plt.title("Top 10 Rules by Occurrences", fontsize=14, fontweight="bold")
        plt.xlabel("Count")
        plt.tight_layout()
        plt.savefig(tmp.name, dpi=110, bbox_inches="tight")
        plt.close()
        return tmp.name
    except Exception as e:
        logger.error(f"Top rules chart failed: {e}")
        return None


def generate_sla_gauge_chart(sla: Dict) -> Optional[str]:
    """Donut/gauge chart of uptime SLA"""
    try:
        if not sla:
            return None
        uptime = float(sla.get("uptime_pct", 0) or 0)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png", dir=UPLOAD_DIR)
        tmp.close()
        fig, ax = plt.subplots(figsize=(5, 5), facecolor="white", subplot_kw=dict(aspect="equal"))
        color = "#10B981" if uptime >= 99.9 else "#F59E0B" if uptime >= 99.0 else "#DC2626"
        ax.pie([uptime, max(100 - uptime, 0)], colors=[color, "#E5E7EB"], startangle=90,
               wedgeprops=dict(width=0.3, edgecolor="white"))
        ax.text(0, 0, f"{uptime:.2f}%", ha="center", va="center", fontsize=26, fontweight="bold")
        ax.text(0, -0.25, "SLA Uptime", ha="center", va="center", fontsize=11, color="#6B7280")
        plt.tight_layout()
        plt.savefig(tmp.name, dpi=110, bbox_inches="tight")
        plt.close()
        return tmp.name
    except Exception as e:
        logger.error(f"SLA gauge chart failed: {e}")
        return None


# ======================== DOCX EXPORT ========================

def generate_docx_report(parsed: Dict, ai_summary: str, period: str = "") -> bytes:
    """Generate DOCX report matching Fasah Weekly Report format"""
    doc = Document()

    # Title
    title = doc.add_heading("FalconOps AI - Weekly Alert Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Period
    if not period:
        now = datetime.now(timezone.utc)
        week_ago = now - timedelta(days=7)
        period = f"{week_ago.strftime('%d %B')} - {now.strftime('%d %B %Y')}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Report Period: {period}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Summary Stats
    doc.add_heading("Alert Summary", level=2)
    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.style = 'Light Grid Accent 1'
    headers = ["Total Alerts", "Critical", "Warning", "Total Occurrences"]
    values = [str(parsed["total_alerts"]), str(parsed["critical_count"]), str(parsed["warning_count"]), str(parsed["total_occurrences"])]
    for i, h in enumerate(headers):
        stats_table.rows[0].cells[i].text = h
    for i, v in enumerate(values):
        stats_table.rows[1].cells[i].text = v

    doc.add_paragraph()

    # AI Summary
    doc.add_heading("AI-Generated Summary", level=2)
    doc.add_paragraph(ai_summary)
    doc.add_paragraph()

    # Alert Details Table
    doc.add_heading("Alert Details", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "Health-Rules Name"
    hdr[1].text = "Critical Alert/Warning"
    hdr[2].text = "Total"
    hdr[3].text = "Summary of Alert Triggered and Action/Issues"

    for alert in parsed["alerts"]:
        row = table.add_row().cells
        row[0].text = alert["rule_name"]
        row[1].text = alert["severity"].capitalize()
        row[2].text = str(alert["count"])
        row[3].text = alert["summary"]

    doc.add_paragraph()
    doc.add_paragraph(f"Report generated by FalconOps AI on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ======================== EXCEL EXPORT ========================

def generate_excel_report(parsed: Dict, ai_summary: str) -> bytes:
    """Generate Excel report with charts"""
    wb = Workbook()

    # Sheet 1: Alert Details
    ws = wb.active
    ws.title = "Alert Details"

    headers = ["Health-Rules Name", "Severity", "Count", "Summary"]
    header_fill = PatternFill(start_color="1F2937", end_color="1F2937", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    thin_border = Border(
        left=Side(style='thin', color='E5E7EB'), right=Side(style='thin', color='E5E7EB'),
        top=Side(style='thin', color='E5E7EB'), bottom=Side(style='thin', color='E5E7EB'),
    )

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border

    for i, alert in enumerate(parsed["alerts"], 2):
        ws.cell(row=i, column=1, value=alert["rule_name"]).border = thin_border
        sev_cell = ws.cell(row=i, column=2, value=alert["severity"].capitalize())
        sev_cell.border = thin_border
        if alert["severity"] == "critical":
            sev_cell.fill = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")
            sev_cell.font = Font(color="DC2626", bold=True)
        ws.cell(row=i, column=3, value=alert["count"]).border = thin_border
        ws.cell(row=i, column=4, value=alert["summary"]).border = thin_border

    ws.column_dimensions['A'].width = 45
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 60

    # Sheet 2: Summary
    ws2 = wb.create_sheet("Summary")
    ws2.cell(row=1, column=1, value="FalconOps AI Weekly Report Summary").font = Font(bold=True, size=14)
    ws2.cell(row=3, column=1, value="Total Alerts:").font = Font(bold=True)
    ws2.cell(row=3, column=2, value=parsed["total_alerts"])
    ws2.cell(row=4, column=1, value="Critical:").font = Font(bold=True)
    ws2.cell(row=4, column=2, value=parsed["critical_count"])
    ws2.cell(row=5, column=1, value="Warning:").font = Font(bold=True)
    ws2.cell(row=5, column=2, value=parsed["warning_count"])
    ws2.cell(row=6, column=1, value="Total Occurrences:").font = Font(bold=True)
    ws2.cell(row=6, column=2, value=parsed["total_occurrences"])
    ws2.cell(row=8, column=1, value="AI Summary:").font = Font(bold=True)
    ws2.cell(row=9, column=1, value=ai_summary)

    # Chart
    if parsed["alerts"]:
        chart = BarChart()
        chart.type = "col"
        chart.title = "Alerts by Count"
        chart.y_axis.title = "Count"
        chart.style = 10

        data_ref = Reference(ws, min_col=3, min_row=1, max_row=min(len(parsed["alerts"]) + 1, 16))
        cats_ref = Reference(ws, min_col=1, min_row=2, max_row=min(len(parsed["alerts"]) + 1, 16))
        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cats_ref)
        chart.shape = 4
        ws2.add_chart(chart, "A12")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ======================== PDF EXPORT (ENTERPRISE) ========================

def _hex_to_color(hex_str: str):
    try:
        h = hex_str.lstrip("#")
        if len(h) == 3:
            h = "".join(c * 2 for c in h)
        return colors.HexColor(f"#{h}")
    except Exception:
        return colors.HexColor("#00E0FF")


def generate_pdf_report(
    parsed: Dict,
    ai_summary: str,
    sla: Dict,
    branding: Dict,
    period: str = "",
    template_sections: Optional[List[Dict]] = None,
) -> bytes:
    """Generate enterprise-grade PDF in Datadog/Splunk style with charts + branding.
    If template_sections is provided, render only those sections in that order;
    otherwise render the full default layout.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.5 * cm, bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    primary = _hex_to_color(branding.get("primary_color", "#00E0FF"))
    secondary = _hex_to_color(branding.get("secondary_color", "#F5B841"))
    dark = colors.HexColor("#0B0E14")

    # Custom styles
    title_style = ParagraphStyle(
        "TitleBig", parent=styles["Title"], fontSize=24, leading=28,
        textColor=dark, alignment=TA_LEFT, spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#6B7280"),
        spaceAfter=14,
    )
    h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontSize=14, textColor=primary,
        spaceBefore=16, spaceAfter=8, fontName="Helvetica-Bold",
    )
    body = ParagraphStyle(
        "Body", parent=styles["BodyText"], fontSize=10, leading=14, textColor=colors.HexColor("#1F2937"),
    )
    footer_style = ParagraphStyle(
        "Footer", parent=styles["Normal"], fontSize=8, textColor=colors.HexColor("#9CA3AF"),
        alignment=TA_CENTER,
    )

    elements = []

    # ============ Template-aware rendering ============
    if template_sections:
        # Prepare period
        if not period:
            now = datetime.now(timezone.utc)
            wk = now - timedelta(days=7)
            period = f"{wk.strftime('%d %b')} – {now.strftime('%d %b %Y')}"

        # Pre-compute charts once
        _sev_chart = generate_severity_chart(parsed)
        _sla_chart = generate_sla_gauge_chart(sla)
        _top_chart = generate_top_rules_chart(parsed)

        def _section_header_logo():
            from .storage_service import get_branding_logo_local_path
            lp = get_branding_logo_local_path(branding.get("logo_url"))
            if lp and os.path.exists(lp):
                try:
                    return [Image(lp, width=3.5 * cm, height=1.2 * cm), Spacer(1, 6)]
                except Exception:
                    return []
            return []

        def _section_title(cfg):
            company = branding.get("company_name", "FalconOps AI")
            custom = (cfg or {}).get("custom_title") or f"{company} · Weekly SOC &amp; AIOps Report"
            return [
                Paragraph(custom, title_style),
                Paragraph(
                    f"Reporting Period: <b>{period}</b> &nbsp;·&nbsp; Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
                    subtitle_style,
                ),
            ]

        def _section_kpi_banner():
            stats_data = [
                ["Risk Posture", "SLA Uptime", "MTTR", "Critical Threats"],
                [
                    sla.get("risk_posture", "N/A"),
                    f"{sla.get('uptime_pct', 0):.2f}%",
                    f"{sla.get('mttr_minutes', 0)} min",
                    str(sla.get("threats_critical", 0)),
                ],
            ]
            risk = sla.get("risk_posture", "Low")
            risk_color = colors.HexColor("#DC2626") if risk == "High" else colors.HexColor("#F59E0B") if risk == "Medium" else colors.HexColor("#10B981")
            t = Table(stats_data, colWidths=[4.3 * cm] * 4)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), dark),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#F9FAFB")),
                ("FONTSIZE", (0, 1), (-1, 1), 14),
                ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
                ("TEXTCOLOR", (0, 1), (0, 1), risk_color),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E5E7EB")),
            ]))
            return [t, Spacer(1, 10)]

        def _section_exec_summary():
            text = (ai_summary or "").replace("\n", "<br/>").replace("**", "")
            return [Paragraph("Executive Summary", h2), Paragraph(text or "No summary available.", body)]

        def _section_sla_table():
            rows = [
                ["Metric", "Value"],
                ["Uptime Compliance", sla.get("sla_compliance", "Unknown")],
                ["Avg Uptime %", f"{sla.get('uptime_pct', 0):.2f}%"],
                ["Monitors Up / Total", f"{sla.get('monitors_up', 0)} / {sla.get('monitors_total', 0)}"],
                ["Monitors Down", str(sla.get("monitors_down", 0))],
                ["Incidents Opened", str(sla.get("incidents_total", 0))],
                ["Incidents Resolved", str(sla.get("incidents_resolved", 0))],
                ["Mean Time to Resolve (MTTR)", f"{sla.get('mttr_minutes', 0)} minutes"],
                ["Total Alerts", str(sla.get("alerts_total", 0))],
                ["Security Threats Detected", str(sla.get("threats_total", 0))],
                ["Critical Threats", str(sla.get("threats_critical", 0))],
            ]
            t = Table(rows, colWidths=[9 * cm, 7.5 * cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), primary),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                ("FONTSIZE", (0, 1), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E5E7EB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F9FAFB"), colors.white]),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            return [Paragraph("SLA &amp; Operations Metrics", h2), t]

        def _section_severity_chart():
            if _sev_chart and os.path.exists(_sev_chart):
                return [Paragraph("Alert Severity", h2), Image(_sev_chart, width=15 * cm, height=8.5 * cm), Spacer(1, 8)]
            return []

        def _section_sla_gauge():
            if _sla_chart and os.path.exists(_sla_chart):
                return [Paragraph("SLA Uptime Gauge", h2), Image(_sla_chart, width=9 * cm, height=9 * cm), Spacer(1, 8)]
            return []

        def _section_top_rules():
            if _top_chart and os.path.exists(_top_chart):
                return [Paragraph("Top 10 Rules", h2), Image(_top_chart, width=16 * cm, height=10 * cm)]
            return []

        def _section_alert_table():
            header = ["Rule Name", "Severity", "Count", "Summary"]
            rows = [header]
            for a in parsed.get("alerts", [])[:20]:
                rows.append([
                    (a.get("rule_name") or "")[:50],
                    (a.get("severity") or "").capitalize(),
                    str(a.get("count", 0)),
                    (a.get("summary") or "")[:80],
                ])
            t = Table(rows, colWidths=[5.5 * cm, 2.3 * cm, 1.7 * cm, 7 * cm], repeatRows=1)
            style = [
                ("BACKGROUND", (0, 0), (-1, 0), dark),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("FONTSIZE", (0, 1), (-1, -1), 8),
                ("ALIGN", (2, 0), (2, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E5E7EB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F9FAFB"), colors.white]),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
            for i, r in enumerate(rows[1:], start=1):
                sev = (r[1] or "").lower()
                if sev == "critical":
                    style.append(("TEXTCOLOR", (1, i), (1, i), colors.HexColor("#DC2626")))
                    style.append(("FONTNAME", (1, i), (1, i), "Helvetica-Bold"))
                elif sev in ("high", "warning"):
                    style.append(("TEXTCOLOR", (1, i), (1, i), colors.HexColor("#F59E0B")))
            t.setStyle(TableStyle(style))
            return [Paragraph("Alert Details (Top 20)", h2), t]

        def _section_custom_text(title, content):
            out = []
            if title:
                out.append(Paragraph(title, h2))
            if content:
                out.append(Paragraph(content.replace("\n", "<br/>"), body))
            return out

        def _section_footer():
            return [Spacer(1, 20), Paragraph(branding.get("footer_text", "Confidential - FalconOps AI SOC Report"), footer_style)]

        def _section_page_break():
            return [PageBreak()]

        for sec in template_sections:
            stype = sec.get("section_type", "")
            cfg = sec.get("config", {})
            if stype == "header_logo":
                elements.extend(_section_header_logo())
            elif stype == "title":
                elements.extend(_section_title(cfg))
            elif stype == "kpi_banner":
                elements.extend(_section_kpi_banner())
            elif stype == "exec_summary":
                elements.extend(_section_exec_summary())
            elif stype == "sla_table":
                elements.extend(_section_sla_table())
            elif stype == "severity_chart":
                elements.extend(_section_severity_chart())
            elif stype == "sla_gauge_chart":
                elements.extend(_section_sla_gauge())
            elif stype == "top_rules_chart":
                elements.extend(_section_top_rules())
            elif stype == "alert_table":
                elements.extend(_section_alert_table())
            elif stype == "custom_text":
                elements.extend(_section_custom_text(sec.get("title") or "", sec.get("content") or ""))
            elif stype == "footer":
                elements.extend(_section_footer())
            elif stype == "page_break":
                elements.extend(_section_page_break())

        def add_page_number_t(canvas, d):
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(colors.HexColor("#9CA3AF"))
            canvas.drawString(1.5 * cm, 1 * cm, branding.get("footer_text", "Confidential"))
            canvas.drawRightString(A4[0] - 1.5 * cm, 1 * cm, f"Page {d.page}")
            canvas.restoreState()

        doc.build(elements, onFirstPage=add_page_number_t, onLaterPages=add_page_number_t)
        return buf.getvalue()

    # ============ Header with Logo (default full-layout path) ============
    from .storage_service import get_branding_logo_local_path
    logo_path = get_branding_logo_local_path(branding.get("logo_url"))
    if logo_path and os.path.exists(logo_path):
        try:
            elements.append(Image(logo_path, width=3.5 * cm, height=1.2 * cm))
        except Exception:
            pass

    elements.append(Spacer(1, 6))
    company = branding.get("company_name", "FalconOps AI")
    elements.append(Paragraph(f"<b>{company}</b> · Weekly SOC &amp; AIOps Report", title_style))

    if not period:
        now = datetime.now(timezone.utc)
        wk = now - timedelta(days=7)
        period = f"{wk.strftime('%d %b')} – {now.strftime('%d %b %Y')}"
    elements.append(Paragraph(
        f"Reporting Period: <b>{period}</b> &nbsp;·&nbsp; Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
        subtitle_style,
    ))

    # ============ Top Stats Banner ============
    stats_data = [
        ["Risk Posture", "SLA Uptime", "MTTR", "Critical Threats"],
        [
            sla.get("risk_posture", "N/A"),
            f"{sla.get('uptime_pct', 0):.2f}%",
            f"{sla.get('mttr_minutes', 0)} min",
            str(sla.get("threats_critical", 0)),
        ],
    ]
    risk = sla.get("risk_posture", "Low")
    risk_color = colors.HexColor("#DC2626") if risk == "High" else colors.HexColor("#F59E0B") if risk == "Medium" else colors.HexColor("#10B981")
    stats_table = Table(stats_data, colWidths=[4.3 * cm] * 4)
    stats_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), dark),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#F9FAFB")),
        ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0, 1), (0, 1), risk_color),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E5E7EB")),
    ]))
    elements.append(stats_table)
    elements.append(Spacer(1, 10))

    # ============ Executive Summary ============
    elements.append(Paragraph("Executive Summary", h2))
    summary_text = (ai_summary or "").replace("\n", "<br/>").replace("**", "")
    elements.append(Paragraph(summary_text or "No summary available.", body))

    # ============ SLA Metrics Table ============
    elements.append(Paragraph("SLA &amp; Operations Metrics", h2))
    sla_rows = [
        ["Metric", "Value"],
        ["Uptime Compliance", sla.get("sla_compliance", "Unknown")],
        ["Avg Uptime %", f"{sla.get('uptime_pct', 0):.2f}%"],
        ["Monitors Up / Total", f"{sla.get('monitors_up', 0)} / {sla.get('monitors_total', 0)}"],
        ["Monitors Down", str(sla.get("monitors_down", 0))],
        ["Incidents Opened", str(sla.get("incidents_total", 0))],
        ["Incidents Resolved", str(sla.get("incidents_resolved", 0))],
        ["Mean Time to Resolve (MTTR)", f"{sla.get('mttr_minutes', 0)} minutes"],
        ["Total Alerts", str(sla.get("alerts_total", 0))],
        ["Security Threats Detected", str(sla.get("threats_total", 0))],
        ["Critical Threats", str(sla.get("threats_critical", 0))],
    ]
    sla_table = Table(sla_rows, colWidths=[9 * cm, 7.5 * cm])
    sla_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), primary),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 10),
        ("ALIGN", (1, 0), (1, -1), "RIGHT"),
        ("FONTSIZE", (0, 1), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E5E7EB")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F9FAFB"), colors.white]),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    elements.append(sla_table)

    # ============ Charts ============
    severity_chart = generate_severity_chart(parsed)
    top_rules_chart = generate_top_rules_chart(parsed)
    sla_chart = generate_sla_gauge_chart(sla)

    if severity_chart or top_rules_chart or sla_chart:
        elements.append(Paragraph("Visual Insights", h2))
        if severity_chart and os.path.exists(severity_chart):
            elements.append(Image(severity_chart, width=15 * cm, height=8.5 * cm))
            elements.append(Spacer(1, 8))
        if sla_chart and os.path.exists(sla_chart):
            elements.append(Image(sla_chart, width=9 * cm, height=9 * cm))
            elements.append(Spacer(1, 8))
        if top_rules_chart and os.path.exists(top_rules_chart):
            elements.append(PageBreak())
            elements.append(Paragraph("Top 10 Rules", h2))
            elements.append(Image(top_rules_chart, width=16 * cm, height=10 * cm))

    # ============ Alert Details Table ============
    elements.append(PageBreak())
    elements.append(Paragraph("Alert Details (Top 20)", h2))
    detail_header = ["Rule Name", "Severity", "Count", "Summary"]
    detail_rows = [detail_header]
    for a in parsed.get("alerts", [])[:20]:
        detail_rows.append([
            (a.get("rule_name") or "")[:50],
            (a.get("severity") or "").capitalize(),
            str(a.get("count", 0)),
            (a.get("summary") or "")[:80],
        ])
    detail_table = Table(detail_rows, colWidths=[5.5 * cm, 2.3 * cm, 1.7 * cm, 7 * cm], repeatRows=1)
    sev_style = [
        ("BACKGROUND", (0, 0), (-1, 0), dark),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("ALIGN", (2, 0), (2, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E5E7EB")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F9FAFB"), colors.white]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    # Color-code severity cells
    for i, row in enumerate(detail_rows[1:], start=1):
        sev = row[1].lower()
        if sev == "critical":
            sev_style.append(("TEXTCOLOR", (1, i), (1, i), colors.HexColor("#DC2626")))
            sev_style.append(("FONTNAME", (1, i), (1, i), "Helvetica-Bold"))
        elif sev in ("high", "warning"):
            sev_style.append(("TEXTCOLOR", (1, i), (1, i), colors.HexColor("#F59E0B")))
    detail_table.setStyle(TableStyle(sev_style))
    elements.append(detail_table)

    # ============ Footer ============
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(
        branding.get("footer_text", "Confidential - FalconOps AI SOC Report"),
        footer_style,
    ))

    # Footer callback for page numbers
    def add_page_number(canvas, d):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#9CA3AF"))
        canvas.drawString(1.5 * cm, 1 * cm, branding.get("footer_text", "Confidential"))
        canvas.drawRightString(A4[0] - 1.5 * cm, 1 * cm, f"Page {d.page}")
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return buf.getvalue()


# ======================== REPORT STORAGE ========================

async def store_report(parsed: Dict, ai_summary: str, docx_bytes: bytes, excel_bytes: bytes, period: str = "", pdf_bytes: Optional[bytes] = None, sla: Optional[Dict] = None, branding: Optional[Dict] = None) -> Dict:
    """Store generated report metadata and files (via storage abstraction)."""
    from .storage_service import save_report_file

    report_id = str(uuid.uuid4())[:12]

    docx_path = save_report_file(report_id, "docx", docx_bytes)
    excel_path = save_report_file(report_id, "xlsx", excel_bytes)
    pdf_path = save_report_file(report_id, "pdf", pdf_bytes) if pdf_bytes else None

    doc = {
        "report_id": report_id,
        "source": parsed.get("source", "manual"),
        "period": period,
        "total_alerts": parsed["total_alerts"],
        "critical_count": parsed["critical_count"],
        "warning_count": parsed["warning_count"],
        "total_occurrences": parsed["total_occurrences"],
        "ai_summary": ai_summary[:2000],
        "alerts": parsed["alerts"][:50],
        "docx_path": docx_path,
        "excel_path": excel_path,
        "pdf_path": pdf_path,
        "sla_metrics": sla or {},
        "branding": {k: v for k, v in (branding or {}).items() if k != "logo_url"},
        "has_pdf": pdf_bytes is not None,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.weekly_reports.insert_one(doc)
    doc.pop("_id", None)
    return doc


async def get_reports(limit: int = 10) -> List[Dict]:
    return await db.weekly_reports.find({}, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)


async def get_report(report_id: str) -> Optional[Dict]:
    return await db.weekly_reports.find_one({"report_id": report_id}, {"_id": 0})
